from docx import Document
from docx.shared import Pt
import os

def create_docx(data: dict) -> str:
    doc = Document()
    
    personal = data.get("personalInfo", {})
    name = personal.get("name", "Name")
    
    # Header
    doc.add_heading(name, 0)
    contact_info = f"{personal.get('email', '')} | {personal.get('phone', '')} | {personal.get('location', '')}"
    p = doc.add_paragraph(contact_info)
    p.alignment = 1 # Center
    
    # Summary
    if data.get("summary"):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(data['summary'])

    # Skills
    if data.get("skills"):
        doc.add_heading('Skills', level=1)
        doc.add_paragraph(data['skills'])

    # Experience
    if data.get("experience"):
        doc.add_heading('Experience', level=1)
        for job in data['experience']:
            p = doc.add_paragraph()
            p.add_run(f"{job['role']}").bold = True
            p.add_run(f" | {job['company']} | {job['dates']}")
            bullets = job['description'].split('\n')
            for bullet in bullets:
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style='List Bullet')

    # Education
    if data.get("education"):
        doc.add_heading('Education', level=1)
        for edu in data['education']:
            doc.add_paragraph(f"{edu['degree']}, {edu['school']} | {edu['dates']}")
            
    # Projects
    if data.get("projects"):
        doc.add_heading('Projects', level=1)
        for proj in data['projects']:
            p = doc.add_paragraph()
            p.add_run(f"{proj['title']}").bold = True
            p.add_run(f" | {proj['technologies']}")
            bullets = proj['description'].split('\n')
            for bullet in bullets:
                if bullet.strip():
                    doc.add_paragraph(bullet.strip(), style='List Bullet')

    # Make it ATS friendly
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    filename = f"{name.replace(' ', '_')}_Resume.docx"
    file_path = os.path.join(os.getcwd(), filename)
    doc.save(file_path)
    return file_path
