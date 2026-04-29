from docx import Document
from docx.shared import Pt

def create_docx(data, name, email, phone):
    doc = Document()

    # Header
    doc.add_heading(name, 0)
    p = doc.add_paragraph(f"{email} | {phone}")
    p.alignment = 1

    # Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(data['summary'])

    # Skills
    doc.add_heading('Skills', level=1)
    doc.add_paragraph(", ".join(data['skills']))

    # Experience
    doc.add_heading('Experience', level=1)
    for job in data['experience']:
        p = doc.add_paragraph()
        p.add_run(f"{job['role']}").bold = True
        p.add_run(f" | {job['company']} | {job['dates']}")
        for bullet in job['bullets']:
            doc.add_paragraph(bullet, style='List Bullet')

    # Education
    doc.add_heading('Education', level=1)
    for edu in data['education']:
        doc.add_paragraph(f"{edu['degree']}, {edu['school']} | {edu['dates']}")

    # Make it ATS friendly
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    doc.save("generated_resume.docx")
    return "generated_resume.docx"